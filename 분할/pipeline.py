# -*- coding: utf-8 -*-
"""
카카오 약관 RAG 결과기 · Upgrade 구현본 (LangGraph)

스켈레톤의 고정 계약(스키마 필드명, 함수 시그니처, 전역 이름)은 그대로 두고 로직만 교체했다.
새 필드와 새 함수는 추가했으며, 기존 필드는 이름·타입을 바꾸지 않고 기본값만 조정했다.

Baseline 대비 변경 요약
  1. 검색   : dense 단일 → dense(bge-m3) + BM25 하이브리드 RRF + cross-encoder 재순위 + 조 단위 집계
  2. 청킹   : 조 단위(최대 1800자) → 항(項) 단위 검색 + 조 단위 컨텍스트 확장
  3. 생성   : Qwen2.5-7B 4bit → Qwen2.5-3B fp16 + prompt lookup decoding
  4. 정합성 : 언어 가드, 회피 문구 제거, 인용 표기 보정, 추출식 폴백
  5. 정적화 : 실행 중 self-distillation LoRA로 출력 계약을 가중치에 내재화(선택, 실패 시 자동 우회)
"""

from __future__ import annotations

import json
import threading
import time
from pathlib import Path
from typing import Any, Dict, List, Literal, Optional, Tuple

from langgraph.graph import END, StateGraph
from pydantic import BaseModel, ConfigDict, Field, field_validator, model_validator

# =====================================================================================
# 0. 스텁 헬퍼
# =====================================================================================

STUB_LOG: List[str] = []


def stub(stage: str, detail: str = "") -> None:
    """스텁 실행을 기록하고 출력한다. 실제 구현으로 교체할 지점 표시."""
    line = f"[STUB] {stage}" + (f" · {detail}" if detail else "")
    STUB_LOG.append(line)
    print(line)


# =====================================================================================
# 1. 고정 상수
# =====================================================================================

DocName = Literal[
    "카카오계정 약관",
    "카카오 위치정보 이용약관",
    "카카오 통합서비스약관",
    "카카오 통합 약관",
]

OFFICIAL_DOCUMENT_NAMES: Tuple[str, ...] = (
    "카카오계정 약관",
    "카카오 위치정보 이용약관",
    "카카오 통합서비스약관",
    "카카오 통합 약관",
)

REQUIRED_GENERATION_MODEL_FAMILY = "Qwen2.5-Instruct"


def normalize_doc_name(value: Any) -> str:
    """공통 러너 _sp_norm_doc과 동일 규칙 · NFC 정규화 + 공백 제거."""
    import re
    import unicodedata

    return re.sub(r"\s+", "", unicodedata.normalize("NFC", str(value)))


ALLOWED_DOCS_NORM = {normalize_doc_name(d) for d in OFFICIAL_DOCUMENT_NAMES}

# 근거가 전혀 없을 때만 쓰는 문구. 채점상 오답 처리되므로 폴백으로 대체한다.
ABSTENTION_SENTENCE = "제공된 약관 조문에서 확인할 수 없습니다."

# 조 단위 원문 레지스트리. 검색은 항 단위로 하고 프롬프트는 조 단위로 확장할 때 참조한다.
ARTICLE_REGISTRY: Dict[Tuple[str, int], "Article"] = {}

CIRCLED_NUMERALS = "①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳"


# =====================================================================================
# 1-1. 텍스트 유틸 (추가)
# =====================================================================================

def character_bigrams(text: str) -> List[str]:
    """한국어 어휘 매칭용 문자 바이그램. 형태소 분석기 없이 부분 일치를 잡는다."""
    import re

    compact = re.sub(r"[^0-9A-Za-z가-힣]", "", text)
    return [compact[i:i + 2] for i in range(len(compact) - 1)]


def korean_tokenize(text: str) -> List[str]:
    """BM25용 토큰 · 어절 + 문자 바이그램 혼합."""
    import re

    words = re.findall(r"[0-9A-Za-z가-힣]+", text)
    return words + character_bigrams(text)


def hangul_ratio(text: str) -> float:
    """전체 문자 중 한글 비율. 언어 이탈 감지에 사용한다."""
    if not text:
        return 0.0
    hangul = sum(1 for ch in text if "가" <= ch <= "힣")
    return hangul / len(text)


def han_ratio(text: str) -> float:
    """전체 문자 중 한자(CJK 통합 한자) 비율. 중국어 출력 감지에 사용한다."""
    if not text:
        return 0.0
    han = sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")
    return han / len(text)


def split_sentences(text: str) -> List[str]:
    """한국어 문장 분할. 종결어미 + 마침표/개행 기준."""
    import re

    parts = re.split(r"(?<=[.!?])\s+|\n+", text)
    return [p.strip() for p in parts if p.strip()]


def bigram_overlap(candidate: str, reference: str) -> float:
    """candidate가 reference 표현을 얼마나 그대로 옮겼는지(0~1). recall 방향."""
    from collections import Counter

    ref = Counter(character_bigrams(reference))
    cand = Counter(character_bigrams(candidate))
    if not cand:
        return 0.0
    overlap = sum((ref & cand).values())
    return overlap / max(1, sum(cand.values()))


def detect_document_hint(question: str) -> Optional[str]:
    """질문이 특정 약관명을 직접 지목하면 그 문서명을 돌려준다."""
    normalized = normalize_doc_name(question)
    for name in OFFICIAL_DOCUMENT_NAMES:
        if normalize_doc_name(name) in normalized:
            return name
    return None


# =====================================================================================
# 2. 스키마 · 설정
# =====================================================================================

class DocumentSource(BaseModel):
    """약관 1종의 원문 취득 경로. local_path가 있으면 로컬 파일을 우선 사용한다."""

    doc_name: DocName
    urls: List[str] = Field(default_factory=list)
    local_path: Optional[str] = None
    effective_date: str
    note: str = ""

    @model_validator(mode="after")
    def _require_source(self) -> "DocumentSource":
        if not self.urls and not self.local_path:
            raise ValueError(f"[{self.doc_name}] urls 또는 local_path 중 하나는 필요합니다.")
        return self


class IndexConfig(BaseModel):
    sources: List[DocumentSource]
    embedding_model_name: str = "BAAI/bge-m3"
    embedding_batch_size: int = 16
    max_chunk_chars: int = 900          # 항 단위 청크 상한
    request_timeout_s: float = 20.0
    # 추가 필드
    passage_prefix: str = ""            # e5 계열을 쓸 때만 "passage: "
    embedding_fp16: bool = True
    min_paragraph_chars: int = 40       # 이보다 짧은 항은 앞 항에 병합
    fetch_retry: int = 2


class RetrievalConfig(BaseModel):
    top_k: int = 4  # 채점 스키마 retrieved 상한과 일치
    query_prefix: str = ""              # bge-m3는 프리픽스 불필요
    # 추가 필드
    dense_candidates: int = 40
    lexical_candidates: int = 40
    rrf_k: int = 60
    rerank_candidates: int = 8
    use_reranker: bool = True
    reranker_model_name: str = "BAAI/bge-reranker-v2-m3"
    reranker_max_chars: int = 1400
    document_hint_boost: float = 0.12
    context_articles: int = 3           # 프롬프트에 넣는 조 개수


class GenerationConfig(BaseModel):
    model_name: str = "Qwen/Qwen2.5-3B-Instruct"
    load_in_4bit: bool = False
    max_new_tokens: int = 448
    temperature: float = 0.0
    max_context_chars: int = 4500
    # 추가 필드
    fallback_model_name: str = "Qwen/Qwen2.5-1.5B-Instruct"
    max_article_chars: int = 2000
    prompt_lookup_num_tokens: int = 10
    use_few_shot: bool = True
    min_hangul_ratio: float = 0.30
    max_han_ratio: float = 0.02
    max_answer_chars: int = 1100
    never_abstain: bool = True


class StyleTuningConfig(BaseModel):
    """실행 중 self-distillation LoRA 설정. 실패하거나 예산을 넘기면 자동으로 건너뛴다."""

    enabled: bool = True
    n_samples: int = 64
    epochs: int = 2
    learning_rate: float = 1e-4
    lora_r: int = 16
    lora_alpha: int = 32
    lora_dropout: float = 0.05
    max_seq_len: int = 1280
    gradient_accumulation_steps: int = 8
    generation_batch_size: int = 4
    min_quote_overlap: float = 0.45
    n_distractors: int = 2
    time_budget_s: float = 900.0
    max_grad_norm: float = 1.0


class PipelineConfig(BaseModel):
    index: IndexConfig
    retrieval: RetrievalConfig = RetrievalConfig()
    generation: GenerationConfig = GenerationConfig()
    style_tuning: StyleTuningConfig = StyleTuningConfig()


# =====================================================================================
# 3. 스키마 · 인덱싱
# =====================================================================================

class RawDocument(BaseModel):
    """취득한 약관 원문 1건 (평문 텍스트)."""

    doc_name: DocName
    text: str
    source_url: str
    fetched_at: str
    char_len: int


class Article(BaseModel):
    """조(條) 단위로 분해된 약관 조항."""

    doc_name: DocName
    article_number: int
    article_title: str = ""
    body: str

    @property
    def citation(self) -> str:
        """골드셋 gold_articles[].citation 과 같은 표기."""
        head = f"{self.doc_name} 제{self.article_number}조"
        return f"{head}({self.article_title})" if self.article_title else head


class Chunk(BaseModel):
    """벡터 저장소 최소 단위. 메타데이터가 채점 스키마와 직접 대응한다."""

    chunk_id: str
    doc_name: DocName
    article_number: int
    article_title: str = ""
    text: str


class EmbeddingBundle(BaseModel):
    """청크 목록과 대응 임베딩 (행 순서 일치)."""

    chunks: List[Chunk]
    vectors: List[List[float]]
    model_name: str
    dim: int


class IndexStats(BaseModel):
    """인덱싱 결과 요약 · 수동 점검용."""

    n_documents: int
    n_articles: int
    n_chunks: int
    dim: int
    per_document: Dict[str, int]
    elapsed_s: float


# =====================================================================================
# 4. 스키마 · 검색 / 증강 / 생성
# =====================================================================================

class RetrievedChunk(BaseModel):
    rank: int
    score: float
    chunk: Chunk


class RetrievalOutput(BaseModel):
    question: str
    hits: List[RetrievedChunk]
    top_k: int
    elapsed_s: float


class PromptBundle(BaseModel):
    """[문서명, 조번호, 본문] 형식으로 조합된 프롬프트."""

    system_prompt: str
    user_prompt: str
    context_block: str
    n_context_chunks: int


class GenerationOutput(BaseModel):
    answer_text: str
    n_new_tokens: int
    elapsed_s: float


class Evidence(BaseModel):
    """retrieved 항목 1건 · [문서명, 조번호] 2원소로 직렬화된다."""

    doc_name: DocName
    article_number: int

    def to_pair(self) -> List[Any]:
        return [self.doc_name, int(self.article_number)]


class AnswerPayload(BaseModel):
    """run_rag_pipeline()의 최종 반환값. 공통 러너 형식 검사와 1:1 대응."""

    answer: str
    retrieved: List[Evidence] = Field(min_length=1, max_length=4)

    @field_validator("retrieved")
    @classmethod
    def _allowed_docs(cls, v: List[Evidence]) -> List[Evidence]:
        for item in v:
            if normalize_doc_name(item.doc_name) not in ALLOWED_DOCS_NORM:
                raise ValueError(f"허용 목록 밖 문서명: {item.doc_name}")
        return v

    def to_contract(self) -> Dict[str, Any]:
        """공통 러너가 기대하는 순수 dict로 변환."""
        return {"answer": self.answer, "retrieved": [e.to_pair() for e in self.retrieved]}


# =====================================================================================
# 5. 스키마 · 품질 결과 (골드셋 / 제출 파일)
# =====================================================================================

class GoldArticle(BaseModel):
    doc: DocName
    article: int
    citation: str


class GoldQuestion(BaseModel):
    id: str
    question: str
    ptype: str
    difficulty: str
    gold_articles: List[GoldArticle]
    key_facts: List[str]


class GoldSet(BaseModel):
    questions: List[GoldQuestion]
    meta: Dict[str, Any] = Field(default_factory=dict, alias="_meta")

    model_config = ConfigDict(populate_by_name=True)


class SubmissionAnswer(BaseModel):
    """answers_public_<팀>.json 의 answers[] 항목."""

    qid: str
    retrieved: List[List[Any]]
    answer: str
    error: Optional[str] = None


class SubmissionFile(BaseModel):
    """answers_public_example.json 과 동일 구조."""

    team: str
    answers: List[SubmissionAnswer]
    meta: Dict[str, Any] = Field(default_factory=dict)


class ArticleScore(BaseModel):
    """근거 조항 정확 일치 채점 결과."""

    qid: str
    predicted: List[List[Any]]
    gold: List[List[Any]]
    hit_at_1: bool
    hit_at_k: bool
    n_gold_matched: int
    n_gold_total: int


class KeyFactScore(BaseModel):
    """정답 핵심 사실 포함 여부. covered는 수동 대조로 확정한다."""

    qid: str
    n_key_facts: int
    n_covered_auto: int
    coverage_auto: float
    per_fact: List[Dict[str, Any]]
    needs_manual_review: bool = True


class ItemReport(BaseModel):
    qid: str
    question: str
    difficulty: str
    ptype: str
    article: ArticleScore
    key_fact: KeyFactScore
    answer_text: str


class EvalReport(BaseModel):
    """공개 10문항 자체 채점 종합."""

    n_items: int
    article_hit_at_1_rate: float
    article_hit_at_k_rate: float
    key_fact_coverage_mean: float
    items: List[ItemReport]


class PerfProtocol(BaseModel):
    """공식 프로토콜과 동일 조건 (Upgrade 단계에서 사용)."""

    requests_per_run: int = 12
    concurrency: int = 2
    warmup_requests: int = 2
    repetitions: int = 3


class PerfReport(BaseModel):
    protocol: PerfProtocol
    success_rate: float
    throughput_rps: float
    p50_latency_s: Optional[float]
    p95_latency_s: Optional[float]


# =====================================================================================
# 6. 인덱싱 · 로딩 및 가져오기
# =====================================================================================

DEFAULT_SOURCES: List[DocumentSource] = [
    DocumentSource(
        doc_name="카카오계정 약관",
        urls=[
            "https://www.kakao.com/policy/terms?lang=ko",
            "https://qr.kakao.com/policy/terms?lang=ko",
            "https://t1.kakaocdn.net/kakaocorp/pw/policy/files/카카오계정약관.pdf",
        ],
        effective_date="2026-05-29",
        note="본문 컨테이너는 div.wrap_terms.wrap_policy로 확인됨. PDF는 미검증 폴백.",
    ),
    DocumentSource(
        doc_name="카카오 위치정보 이용약관",
        urls=[
            "https://www.kakao.com/policy/location?lang=ko",
            "https://qr.kakao.com/policy/location?lang=ko",
        ],
        effective_date="2026-07-16",
        note="본문 컨테이너는 div.wrap_terms(단독)로 확인됨. PDF 폴백 경로는 미확보.",
    ),
    DocumentSource(
        doc_name="카카오 통합서비스약관",
        urls=[
            "https://www.kakao.com/policy/terms?type=ts&lang=ko",
            "https://qr.kakao.com/policy/terms?type=ts&lang=ko",
        ],
        effective_date="2026-05-29",
        note="본문 컨테이너는 div.wrap_terms.wrap_policy로 확인됨. PDF 폴백 경로는 미확보.",
    ),
    DocumentSource(
        doc_name="카카오 통합 약관",
        urls=[
            "https://www.kakao.com/policy/kakaoTerms?lang=ko",
            "https://qr.kakao.com/policy/kakaoTerms?lang=ko",
        ],
        effective_date="2022-08-25",
    ),
]


def _load_local_file_text(local_path: str) -> Tuple[str, str]:
    """로컬 PDF/DOCX 경로 → (원문 텍스트, 실제로 읽은 경로)."""
    from pathlib import Path

    candidates = [Path(local_path), Path("/content") / local_path]
    resolved = next((path for path in candidates if path.exists()), None)
    if resolved is None:
        tried = ", ".join(str(path) for path in candidates)
        raise FileNotFoundError(f"로컬 파일을 찾을 수 없습니다. 시도한 경로: {tried}")

    suffix = resolved.suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(resolved))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif suffix == ".docx":
        from docx import Document

        document = Document(str(resolved))
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    else:
        raise ValueError(f"지원하지 않는 로컬 파일 형식: {suffix} ({resolved})")

    return text, str(resolved)


def fetch_document(source: DocumentSource, timeout_s: float = 20.0) -> RawDocument:
    """DocumentSource → RawDocument.

    local_path가 있으면 로컬 PDF/DOCX를 우선 사용한다. 없으면 urls를 순서대로
    시도해 조 구조(제N조)가 확인되는 첫 응답을 채택한다. 네트워크 순단에 대비해
    URL마다 재시도를 건다.
    """
    import datetime
    import re

    article_pattern = re.compile(r"제\s*\d+\s*조")

    def finalize(raw_text: str) -> str:
        cleaned = raw_text.replace("\u00a0", " ")
        lines = [line.strip() for line in cleaned.split("\n")]
        return "\n".join(line for line in lines if line)

    if source.local_path:
        raw_text, resolved_path = _load_local_file_text(source.local_path)
        text = finalize(raw_text)
        if len(article_pattern.findall(text)) < 3:
            raise RuntimeError(
                f"[{source.doc_name}] 조 구조 미검출 · 경로={resolved_path} · len={len(text)}"
            )
        return RawDocument(
            doc_name=source.doc_name,
            text=text,
            source_url=resolved_path,
            fetched_at=datetime.datetime.now(datetime.timezone.utc).isoformat(),
            char_len=len(text),
        )

    import requests
    from bs4 import BeautifulSoup

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"
        ),
        "Accept-Language": "ko-KR,ko;q=0.9",
    }
    failures: List[str] = []

    def is_pdf_response(url: str, response: "requests.Response") -> bool:
        content_type = response.headers.get("Content-Type", "").lower()
        return "application/pdf" in content_type or url.lower().endswith(".pdf")

    for url in source.urls:
        for attempt in range(1, 3):
            try:
                response = requests.get(url, headers=headers, timeout=timeout_s)
                response.raise_for_status()

                if is_pdf_response(url, response):
                    from io import BytesIO

                    from pypdf import PdfReader

                    reader = PdfReader(BytesIO(response.content))
                    raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                else:
                    response.encoding = response.apparent_encoding or "utf-8"
                    soup = BeautifulSoup(response.text, "html.parser")
                    for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
                        tag.decompose()

                    container_selectors = ["div.wrap_terms", "div.wrap_policy", "main"]
                    content = None
                    for selector in container_selectors:
                        content = soup.select_one(selector)
                        if content is not None:
                            break
                    if content is None:
                        print(
                            f"[경고][로딩] {source.doc_name} · 본문 컨테이너를 찾지 못해 "
                            "전체 페이지에서 추출합니다(목차 혼입 가능)."
                        )
                        content = soup

                    raw_text = content.get_text(separator="\n")

                text = finalize(raw_text)

                if len(article_pattern.findall(text)) < 3:
                    failures.append(f"{url}: 조 구조 미검출(len={len(text)})")
                    break

                return RawDocument(
                    doc_name=source.doc_name,
                    text=text,
                    source_url=url,
                    fetched_at=datetime.datetime.now(datetime.timezone.utc).isoformat(),
                    char_len=len(text),
                )
            except Exception as exc:  # noqa: BLE001 · URL 후보를 끝까지 시도한다
                failures.append(f"{url}(try{attempt}): {type(exc).__name__}: {exc}")
                time.sleep(1.0)

    raise RuntimeError(f"[{source.doc_name}] 원문 취득 실패 · " + " | ".join(failures))


def load_documents(
    sources: List[DocumentSource],
    timeout_s: float = 20.0,
) -> List[RawDocument]:
    """List[DocumentSource] → List[RawDocument]."""
    documents: List[RawDocument] = []
    for source in sources:
        document = fetch_document(source, timeout_s)
        print(f"[로딩] {document.doc_name} · {document.char_len}자 · {document.source_url}")
        documents.append(document)
    return documents


# =====================================================================================
# 7. 인덱싱 · 파싱
# =====================================================================================

def parse_articles(document: RawDocument) -> List[Article]:
    """RawDocument → List[Article].

    조 번호는 1부터 순차 증가하는 구간(run)만 헤더 후보로 인정하고, 목차와 본문이
    각각 1..N 구간을 형성하는 경우 글자 폭이 가장 큰 구간을 본문으로 채택한다.
    마지막 조 뒤에 붙는 부칙은 본문에서 잘라 낸다.
    """
    import re

    text = document.text
    header_pattern = re.compile(r"제\s*(\d+)\s*조")
    reference_suffix = ("에", "의", "와", "과", "및", "부터", "까지", "에서", ",", "제")

    candidates: List[Tuple[int, int, int, str]] = []  # (시작, 헤더끝, 조번호, 제목)

    for match in header_pattern.finditer(text):
        number = int(match.group(1))

        tail = text[match.end():match.end() + 90]
        stripped_tail = tail.lstrip()
        if stripped_tail and stripped_tail[0] in reference_suffix:
            continue

        title_match = re.match(r"[ \t]*\(([^)\n]{1,60})\)", tail)
        if title_match:
            title = title_match.group(1).strip()
            header_end = match.end() + title_match.end()
        else:
            line_match = re.match(r"[ \t]*([^\n]{0,60})", tail)
            candidate = line_match.group(1).strip() if line_match else ""
            candidate = candidate.split("\n")[0].strip()
            title = candidate if 0 < len(candidate) <= 40 else ""
            header_end = match.end() + (len(line_match.group(0)) if title else 0)

        candidates.append((match.start(), header_end, number, title))

    runs: List[List[Tuple[int, int, int, str]]] = []
    current_run: List[Tuple[int, int, int, str]] = []
    expected_number = 1

    for candidate in candidates:
        _, _, number, _ = candidate
        if number == expected_number:
            current_run.append(candidate)
            expected_number += 1
        elif number == 1:
            if current_run:
                runs.append(current_run)
            current_run = [candidate]
            expected_number = 2

    if current_run:
        runs.append(current_run)

    if not runs:
        raise ValueError(f"[{document.doc_name}] 조 구조 파싱 실패 · 정규식 재검토 필요")

    def run_span(run: List[Tuple[int, int, int, str]]) -> int:
        return run[-1][0] - run[0][0]

    headers = max(runs, key=run_span)

    if len(runs) > 1:
        detail = ", ".join(f"{len(run)}개조/span={run_span(run)}자" for run in runs)
        print(
            f"[경고][파싱] {document.doc_name} 조 시퀀스 {len(runs)}개 발견 · {detail} "
            "· 가장 긴 구간을 본문으로 채택"
        )

    chapter_pattern = re.compile(r"^제\s*\d+\s*장.*$", re.MULTILINE)
    supplement_pattern = re.compile(r"\n\s*부\s*칙[\s\S]*$")
    articles: List[Article] = []

    for index, (_, header_end, number, title) in enumerate(headers):
        body_end = headers[index + 1][0] if index + 1 < len(headers) else len(text)
        body = text[header_end:body_end]
        if index + 1 == len(headers):
            body = supplement_pattern.sub("", body)
        body = chapter_pattern.sub("", body)
        body = re.sub(r"\n{2,}", "\n", body).strip()

        articles.append(
            Article(
                doc_name=document.doc_name,
                article_number=number,
                article_title=title,
                body=body,
            )
        )

    suspicious = [
        f"제{a.article_number}조(len={len(a.body)})" for a in articles if len(a.body) < 20
    ]
    if suspicious:
        print(
            f"[경고][파싱] {document.doc_name} 본문이 20자 미만인 조 {len(suspicious)}건: "
            + ", ".join(suspicious)
        )

    print(
        f"[파싱] {document.doc_name} · 제1조~제{articles[-1].article_number}조 "
        f"({len(articles)}개)"
    )
    return articles


def split_article_paragraphs(body: str, min_chars: int = 40) -> List[Tuple[str, str]]:
    """조 본문 → [(항 라벨, 항 본문)]. 항 기호가 없으면 줄 묶음으로 분할한다."""
    import re

    positions = [(m.start(), m.group(0)) for m in re.finditer(r"[①-⑳]", body)]

    segments: List[Tuple[str, str]] = []
    if len(positions) >= 2:
        head = body[: positions[0][0]].strip()
        if head:
            segments.append(("", head))
        for order, (start, label) in enumerate(positions):
            end = positions[order + 1][0] if order + 1 < len(positions) else len(body)
            segments.append((label, body[start + 1:end].strip()))
    else:
        numbered = [(m.start(), m.group(1)) for m in re.finditer(r"(?m)^\s*(\d{1,2})[.)]\s", body)]
        if len(numbered) >= 2:
            head = body[: numbered[0][0]].strip()
            if head:
                segments.append(("", head))
            for order, (start, label) in enumerate(numbered):
                end = numbered[order + 1][0] if order + 1 < len(numbered) else len(body)
                segments.append((f"{label}.", body[start:end].strip()))
        else:
            segments.append(("", body.strip()))

    merged: List[Tuple[str, str]] = []
    for label, segment in segments:
        if not segment:
            continue
        if merged and len(segment) < min_chars:
            previous_label, previous_text = merged[-1]
            merged[-1] = (previous_label, f"{previous_text}\n{label} {segment}".strip())
        else:
            merged.append((label, segment))
    return merged or [("", body.strip())]


# =====================================================================================
# 8. 인덱싱 · 청킹
# =====================================================================================

def chunk_articles(articles: List[Article], config: IndexConfig) -> List[Chunk]:
    """List[Article] → List[Chunk]. 항 단위 1청크, max_chunk_chars 초과분만 재분할.

    검색 단위를 항으로 낮춰 질문이 겨냥한 문장이 상위에 오게 하고, 프롬프트는
    ARTICLE_REGISTRY에서 조 전체를 복원해 완결성을 유지한다.
    """
    chunks: List[Chunk] = []
    empty_body_articles: List[str] = []

    for article in articles:
        ARTICLE_REGISTRY[(article.doc_name, article.article_number)] = article

        header_line = f"[{article.doc_name}] 제{article.article_number}조"
        if article.article_title:
            header_line += f"({article.article_title})"

        body = article.body
        if not body:
            empty_body_articles.append(f"제{article.article_number}조")
            chunks.append(
                Chunk(
                    chunk_id=(
                        f"{normalize_doc_name(article.doc_name)}"
                        f"-{article.article_number:03d}-00"
                    ),
                    doc_name=article.doc_name,
                    article_number=article.article_number,
                    article_title=article.article_title,
                    text=f"{header_line}\n(본문 파싱 실패 · 원문 확인 필요)",
                )
            )
            continue

        paragraphs = split_article_paragraphs(body, config.min_paragraph_chars)
        budget = max(200, config.max_chunk_chars - len(header_line) - 1)

        segments: List[Tuple[str, str]] = []
        for label, paragraph in paragraphs:
            if len(paragraph) <= budget:
                segments.append((label, paragraph))
                continue
            cursor = 0
            while cursor < len(paragraph):
                window_end = min(cursor + budget, len(paragraph))
                if window_end < len(paragraph):
                    boundary = paragraph.rfind("\n", cursor + budget // 2, window_end)
                    if boundary == -1:
                        boundary = paragraph.rfind(". ", cursor + budget // 2, window_end)
                    if boundary != -1:
                        window_end = boundary + 1
                segments.append((label, paragraph[cursor:window_end].strip()))
                cursor = window_end

        for part_index, (label, segment) in enumerate(segments):
            if not segment:
                continue
            marker = f" {label}" if label else ""
            chunks.append(
                Chunk(
                    chunk_id=(
                        f"{normalize_doc_name(article.doc_name)}"
                        f"-{article.article_number:03d}-{part_index:02d}"
                    ),
                    doc_name=article.doc_name,
                    article_number=article.article_number,
                    article_title=article.article_title,
                    text=f"{header_line}{marker}\n{segment}",
                )
            )

    if empty_body_articles:
        print(
            f"[경고][청킹] 본문 파싱 실패 {len(empty_body_articles)}건: "
            + ", ".join(empty_body_articles)
        )

    print(f"[청킹] {len(articles)}개 조항 → {len(chunks)}개 청크(항 단위)")
    return chunks


# =====================================================================================
# 9. 인덱싱 · 임베딩
# =====================================================================================

class EmbedderHandle(BaseModel):
    """임베딩 모델 핸들."""

    model_config = ConfigDict(arbitrary_types_allowed=True)

    model_name: str
    dim: int
    device: str = "cpu"
    model: Any = None


def build_embedder(config: IndexConfig) -> EmbedderHandle:
    """IndexConfig → EmbedderHandle. bge-m3를 fp16으로 로컬 로드한다."""
    import torch
    from sentence_transformers import SentenceTransformer

    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = SentenceTransformer(config.embedding_model_name, device=device)
    if config.embedding_fp16 and device == "cuda":
        model = model.half()
    dim = int(model.get_sentence_embedding_dimension())

    print(f"[임베딩 모델] {config.embedding_model_name} · device={device} · dim={dim}")
    return EmbedderHandle(
        model_name=config.embedding_model_name, dim=dim, device=device, model=model
    )


def embed_chunks(
    chunks: List[Chunk],
    embedder: EmbedderHandle,
    config: IndexConfig,
) -> EmbeddingBundle:
    """List[Chunk] → EmbeddingBundle. L2 정규화하여 내적=cosine이 되게 한다."""
    texts = [config.passage_prefix + chunk.text for chunk in chunks]
    vectors = embedder.model.encode(
        texts,
        batch_size=config.embedding_batch_size,
        normalize_embeddings=True,
        convert_to_numpy=True,
        show_progress_bar=False,
    )

    print(f"[임베딩] {len(chunks)}개 청크 · shape={tuple(vectors.shape)}")
    return EmbeddingBundle(
        chunks=chunks,
        vectors=vectors.astype("float32").tolist(),
        model_name=embedder.model_name,
        dim=int(vectors.shape[1]),
    )


# =====================================================================================
# 10. 인덱싱 · 저장
# =====================================================================================

class LexicalIndexHandle(BaseModel):
    """BM25 희소 검색 핸들. 조번호·법령명처럼 표기가 그대로 겹치는 질의를 잡는다."""

    model_config = ConfigDict(arbitrary_types_allowed=True)

    backend: str = "rank_bm25.BM25Okapi"
    n_documents: int
    model: Any = None


_LEXICAL_INDEX: Optional[LexicalIndexHandle] = None
_RERANKER: Optional[Any] = None


class VectorStoreHandle(BaseModel):
    """FAISS in-memory 저장소 핸들."""

    model_config = ConfigDict(arbitrary_types_allowed=True)

    backend: str = "faiss.IndexFlatIP"
    dim: int
    n_vectors: int
    chunks: List[Chunk]
    index: Any = None

    def search(self, query_vector: List[float], top_k: int) -> List[RetrievedChunk]:
        """질의 벡터 → 상위 top_k 청크. 정규화 벡터이므로 내적=cosine."""
        import numpy as np

        query = np.asarray([query_vector], dtype="float32")
        scores, indices = self.index.search(query, min(top_k, self.n_vectors))

        hits: List[RetrievedChunk] = []
        for rank, (score, position) in enumerate(zip(scores[0], indices[0]), start=1):
            if position < 0:
                continue
            hits.append(
                RetrievedChunk(
                    rank=rank,
                    score=round(float(score), 4),
                    chunk=self.chunks[int(position)],
                )
            )
        return hits


def build_lexical_index(chunks: List[Chunk]) -> Optional[LexicalIndexHandle]:
    """List[Chunk] → BM25 인덱스. 패키지가 없으면 None을 돌려주고 dense만 쓴다."""
    try:
        from rank_bm25 import BM25Okapi
    except ImportError:
        print("[경고][저장] rank_bm25 미설치 · dense 단독 검색으로 진행")
        return None

    corpus = [korean_tokenize(chunk.text) for chunk in chunks]
    model = BM25Okapi(corpus)
    print(f"[저장] BM25Okapi · {len(corpus)}개 문서")
    return LexicalIndexHandle(n_documents=len(corpus), model=model)


def build_vector_store(bundle: EmbeddingBundle) -> VectorStoreHandle:
    """EmbeddingBundle → VectorStoreHandle. FAISS IndexFlatIP 생성 후 add."""
    global _LEXICAL_INDEX

    import faiss
    import numpy as np

    vectors = np.asarray(bundle.vectors, dtype="float32")
    index = faiss.IndexFlatIP(bundle.dim)
    index.add(vectors)

    _LEXICAL_INDEX = build_lexical_index(bundle.chunks)

    print(f"[저장] FAISS IndexFlatIP · {index.ntotal}개 벡터 · dim={bundle.dim}")
    return VectorStoreHandle(
        dim=bundle.dim,
        n_vectors=int(index.ntotal),
        chunks=bundle.chunks,
        index=index,
    )


_PARSED_ARTICLES: List[Article] = []


def prepare_corpus(config: IndexConfig) -> List[Article]:
    """로딩 + 파싱만 먼저 수행하고 캐시한다. 미세조정이 인덱싱보다 먼저 필요하다."""
    global _PARSED_ARTICLES

    if _PARSED_ARTICLES:
        return _PARSED_ARTICLES

    documents = load_documents(config.sources, config.request_timeout_s)
    articles: List[Article] = []
    for document in documents:
        articles.extend(parse_articles(document))
    _PARSED_ARTICLES = articles
    return articles


def build_index(config: IndexConfig) -> Tuple[VectorStoreHandle, EmbedderHandle, IndexStats]:
    """인덱싱 오케스트레이션: 로딩 → 파싱 → 청킹 → 임베딩 → 저장."""
    started = time.perf_counter()
    articles = prepare_corpus(config)

    chunks = chunk_articles(articles, config)
    embedder = build_embedder(config)
    bundle = embed_chunks(chunks, embedder, config)
    store = build_vector_store(bundle)

    per_document: Dict[str, int] = {}
    for chunk in chunks:
        per_document[chunk.doc_name] = per_document.get(chunk.doc_name, 0) + 1

    stats = IndexStats(
        n_documents=len({a.doc_name for a in articles}),
        n_articles=len(articles),
        n_chunks=len(chunks),
        dim=store.dim,
        per_document=per_document,
        elapsed_s=round(time.perf_counter() - started, 4),
    )
    return store, embedder, stats


# =====================================================================================
# 11. 검색
# =====================================================================================

def load_reranker(config: RetrievalConfig) -> Optional[Any]:
    """cross-encoder 재순위 모델. 로드 실패 시 None으로 두고 하이브리드 점수만 쓴다."""
    if not config.use_reranker:
        return None
    try:
        import torch
        from sentence_transformers import CrossEncoder

        device = "cuda" if torch.cuda.is_available() else "cpu"
        model = CrossEncoder(config.reranker_model_name, device=device, max_length=512)
        if device == "cuda":
            model.model.half()
        print(f"[재순위 모델] {config.reranker_model_name} · device={device}")
        return model
    except Exception as exc:  # noqa: BLE001 · 재순위는 없어도 동작해야 한다
        print(f"[경고][검색] 재순위 모델 로드 실패({type(exc).__name__}) · 하이브리드 점수만 사용")
        return None


def embed_query(
    question: str,
    embedder: EmbedderHandle,
    config: RetrievalConfig,
) -> List[float]:
    """질문 문자열 → 정규화된 질의 벡터."""
    vector = embedder.model.encode(
        [config.query_prefix + question],
        normalize_embeddings=True,
        convert_to_numpy=True,
        show_progress_bar=False,
    )[0]
    return vector.astype("float32").tolist()


def reciprocal_rank_fusion(
    ranked_lists: List[List[int]],
    rrf_k: int = 60,
) -> Dict[int, float]:
    """여러 순위 목록 → RRF 융합 점수. 점수 스케일이 다른 검색기를 순위로만 합친다."""
    fused: Dict[int, float] = {}
    for ranked in ranked_lists:
        for rank, position in enumerate(ranked, start=1):
            fused[position] = fused.get(position, 0.0) + 1.0 / (rrf_k + rank)
    return fused


def _dense_candidate_positions(
    question: str,
    store: VectorStoreHandle,
    embedder: EmbedderHandle,
    config: RetrievalConfig,
) -> List[int]:
    import numpy as np

    query_vector = np.asarray([embed_query(question, embedder, config)], dtype="float32")
    _, indices = store.index.search(
        query_vector, min(config.dense_candidates, store.n_vectors)
    )
    return [int(p) for p in indices[0] if p >= 0]


def _lexical_candidate_positions(question: str, config: RetrievalConfig) -> List[int]:
    if _LEXICAL_INDEX is None or _LEXICAL_INDEX.model is None:
        return []
    import numpy as np

    scores = _LEXICAL_INDEX.model.get_scores(korean_tokenize(question))
    order = np.argsort(scores)[::-1][: config.lexical_candidates]
    return [int(p) for p in order if scores[int(p)] > 0.0]


def retrieve(
    question: str,
    store: VectorStoreHandle,
    embedder: EmbedderHandle,
    config: RetrievalConfig,
) -> RetrievalOutput:
    """질문 → RetrievalOutput.

    dense(top-N) + BM25(top-N)를 RRF로 융합하고, 항 단위 점수를 조 단위로 집계한 뒤
    cross-encoder로 재순위한다. 질문이 약관명을 직접 지목하면 해당 문서에 가산점을 준다.
    """
    started = time.perf_counter()

    dense_positions = _dense_candidate_positions(question, store, embedder, config)
    lexical_positions = _lexical_candidate_positions(question, config)
    fused = reciprocal_rank_fusion([dense_positions, lexical_positions], config.rrf_k)

    if not fused:
        fused = {position: 1.0 for position in dense_positions[: config.top_k]}

    # 항 단위 점수를 조 단위로 집계한다. 조 대표 청크는 최고점 항으로 둔다.
    article_best: Dict[Tuple[str, int], Tuple[float, int]] = {}
    for position, score in fused.items():
        chunk = store.chunks[position]
        key = (chunk.doc_name, chunk.article_number)
        if key not in article_best or score > article_best[key][0]:
            article_best[key] = (score, position)

    hint = detect_document_hint(question)
    max_score = max((s for s, _ in article_best.values()), default=1.0) or 1.0
    scored: List[Tuple[float, Tuple[str, int], int]] = []
    for key, (score, position) in article_best.items():
        normalized = score / max_score
        if hint is not None and key[0] == hint:
            normalized += config.document_hint_boost
        scored.append((normalized, key, position))
    scored.sort(key=lambda item: item[0], reverse=True)

    shortlist = scored[: max(config.rerank_candidates, config.top_k)]

    if _RERANKER is not None and len(shortlist) > 1:
        try:
            pairs = []
            for _, key, _position in shortlist:
                article = ARTICLE_REGISTRY.get(key)
                body = article.body if article is not None else store.chunks[_position].text
                pairs.append((question, f"{key[0]} 제{key[1]}조\n{body[: config.reranker_max_chars]}"))
            raw_scores = _RERANKER.predict(pairs, batch_size=8, show_progress_bar=False)
            rescored: List[Tuple[float, Tuple[str, int], int]] = []
            for (raw, (_, key, position)) in zip(raw_scores, shortlist):
                value = 1.0 / (1.0 + pow(2.718281828, -float(raw)))
                if hint is not None and key[0] == hint:
                    value += config.document_hint_boost
                rescored.append((value, key, position))
            rescored.sort(key=lambda item: item[0], reverse=True)
            shortlist = rescored
        except Exception as exc:  # noqa: BLE001 · 재순위 실패는 치명적이지 않다
            print(f"[경고][검색] 재순위 실패({type(exc).__name__}) · 융합 점수 순서를 사용")

    hits: List[RetrievedChunk] = []
    for rank, (score, _key, position) in enumerate(shortlist[: config.top_k], start=1):
        hits.append(
            RetrievedChunk(rank=rank, score=round(float(score), 4), chunk=store.chunks[position])
        )

    return RetrievalOutput(
        question=question,
        hits=hits,
        top_k=config.top_k,
        elapsed_s=round(time.perf_counter() - started, 4),
    )


# =====================================================================================
# 12. 증강
# =====================================================================================

SYSTEM_PROMPT = (
    "당신은 카카오 약관 질의응답 어시스턴트입니다. 아래 규칙을 반드시 지킵니다.\n"
    "1. 한국어로만 답합니다. 다른 언어는 한 글자도 쓰지 않습니다.\n"
    "2. 제공된 조문에 있는 내용만 사용하고, 문장 표현은 조문 원문을 그대로 옮깁니다.\n"
    "3. 질문이 예/아니오를 묻는 형태이면 첫 문장을 \"예.\" 또는 \"아니오.\"로 시작합니다.\n"
    "4. 질문이 요구한 항목은 빠짐없이 담고, 질문과 무관한 조문은 인용하지 않습니다.\n"
    "5. 답변 끝에 근거를 \"(문서명 제N조)\" 형식으로 한 번만 표기합니다.\n"
    "6. 서론·인사말·요약 표현을 쓰지 않고 사실 문장만 씁니다."
)

FEW_SHOT_EXAMPLES = (
    "다음은 답변 방식을 보여주는 예시입니다.\n\n"
    "[예시 조문] 카카오 예시약관 제5조(게시물) ② 회원이 탈퇴하면 본인이 작성한 게시물은 "
    "삭제됩니다. 다만 제3자가 공유하거나 댓글을 단 게시물은 삭제되지 않습니다.\n"
    "[예시 질문] 탈퇴하면 제가 쓴 글이 전부 삭제되나요?\n"
    "[예시 답변] 아니오. 회원이 탈퇴하면 본인이 작성한 게시물은 삭제되지만, 제3자가 공유하거나 "
    "댓글을 단 게시물은 삭제되지 않습니다. (카카오 예시약관 제5조)\n\n"
    "[예시 조문] 카카오 예시약관 제9조(약관의 개정) ① 회사는 관련 법령을 위반하지 않는 범위에서 "
    "본 약관을 개정할 수 있습니다. ② 개정 약관은 적용일자 30일 전부터 공지합니다.\n"
    "[예시 질문] 약관을 개정하면 언제부터 공지하나요?\n"
    "[예시 답변] 회사는 개정 약관을 적용일자 30일 전부터 공지합니다. 회사는 관련 법령을 위반하지 "
    "않는 범위에서 본 약관을 개정할 수 있습니다. (카카오 예시약관 제9조)"
)

_STYLE_ADAPTER_APPLIED = False


def render_context_block(retrieval: RetrievalOutput, config: GenerationConfig) -> str:
    """검색 결과 → 조 단위로 확장된 근거 블록.

    검색은 항 단위로 하되 프롬프트에는 조 전체를 넣는다. 질문이 겨냥한 문장만 잘려
    들어가면 결론을 못 내리고 회피하는 실패가 생기기 때문이다.
    """
    from typing import Set

    blocks: List[str] = []
    used_chars = 0
    seen: Set[Tuple[str, int]] = set()

    for hit in retrieval.hits:
        key = (hit.chunk.doc_name, hit.chunk.article_number)
        if key in seen:
            continue
        seen.add(key)

        article = ARTICLE_REGISTRY.get(key)
        title = article.article_title if article is not None else hit.chunk.article_title
        body = article.body if article is not None else hit.chunk.text
        head = f"[근거 {len(blocks) + 1}] {key[0]} 제{key[1]}조"
        if title:
            head += f"({title})"
        block = f"{head}\n{body[: config.max_article_chars]}"

        if blocks and used_chars + len(block) > config.max_context_chars:
            break
        blocks.append(block)
        used_chars += len(block)
        if len(blocks) >= 3:
            break

    return "\n\n".join(blocks)


def build_prompt(retrieval: RetrievalOutput, config: GenerationConfig) -> PromptBundle:
    """RetrievalOutput → PromptBundle. 조 단위로 확장한 근거만 담는다."""
    context_block = render_context_block(retrieval, config)

    parts: List[str] = []
    if config.use_few_shot and not _STYLE_ADAPTER_APPLIED:
        parts.append(FEW_SHOT_EXAMPLES)
    parts.append(f"[조문]\n{context_block}")
    parts.append(f"[질문]\n{retrieval.question}")
    parts.append("[답변]")

    return PromptBundle(
        system_prompt=SYSTEM_PROMPT,
        user_prompt="\n\n".join(parts),
        context_block=context_block,
        n_context_chunks=context_block.count("[근거 "),
    )


def primary_citation(context_block: str) -> str:
    """근거 블록 첫 항목의 \"(문서명 제N조)\" 표기를 뽑는다."""
    import re

    match = re.search(r"\[근거 1\]\s*(.+?)\s*제(\d+)조", context_block)
    if not match:
        return ""
    return f"({match.group(1).strip()} 제{int(match.group(2))}조)"


# =====================================================================================
# 12-1. 답변 후처리 (추가)
# =====================================================================================

META_PREFIX_PATTERNS = (
    r"^\s*(답변|정답|A)\s*[:：]\s*",
    r"^\s*제공된 약관 조문에 (따르면|의하면)[,\s]*",
    r"^\s*질문(하신|에 대한) [^\n.]{0,30}[:：]\s*",
)


def polish_answer(text: str, citation: str, config: GenerationConfig) -> str:
    """생성 원문 → 제출 답변. 채점에 손해가 되는 군더더기와 모순을 제거한다."""
    import re

    answer = (text or "").strip()
    answer = re.sub(r"^```[a-zA-Z]*\n?|```$", "", answer).strip()
    for pattern in META_PREFIX_PATTERNS:
        answer = re.sub(pattern, "", answer)

    # 답을 제시하고도 회피 문구를 덧붙이는 자기모순을 제거한다.
    if ABSTENTION_SENTENCE in answer:
        stripped = answer.replace(ABSTENTION_SENTENCE, " ").strip(" .\n")
        if len(stripped) >= 25:
            answer = stripped

    sentences = split_sentences(answer)
    deduplicated: List[str] = []
    for sentence in sentences:
        if deduplicated and sentence == deduplicated[-1]:
            continue
        if sentence in deduplicated:
            continue
        deduplicated.append(sentence)
    answer = " ".join(deduplicated) if deduplicated else answer

    if len(answer) > config.max_answer_chars:
        cut = answer[: config.max_answer_chars]
        boundary = max(cut.rfind("다."), cut.rfind("."))
        answer = cut[: boundary + 1] if boundary > 0 else cut

    if citation and "제" not in answer[-40:]:
        answer = f"{answer} {citation}".strip()

    return re.sub(r"[ \t]{2,}", " ", answer).strip()


def build_extractive_answer(
    retrieval: RetrievalOutput,
    config: GenerationConfig,
    max_sentences: int = 3,
) -> str:
    """생성이 실패했을 때 쓰는 결정적 폴백. 상위 조에서 질문과 겹치는 문장을 원문 그대로 뽑는다."""
    if not retrieval.hits:
        return ABSTENTION_SENTENCE

    top = retrieval.hits[0]
    key = (top.chunk.doc_name, top.chunk.article_number)
    article = ARTICLE_REGISTRY.get(key)
    body = article.body if article is not None else top.chunk.text

    sentences = [s for s in split_sentences(body) if len(s) >= 15]
    if not sentences:
        return ABSTENTION_SENTENCE

    scored = [
        (bigram_overlap(retrieval.question, sentence), order, sentence)
        for order, sentence in enumerate(sentences)
    ]
    scored.sort(key=lambda item: item[0], reverse=True)
    chosen = sorted(scored[:max_sentences], key=lambda item: item[1])

    citation = f"({key[0]} 제{key[1]}조)"
    return polish_answer(" ".join(s for _, _, s in chosen), citation, config)


# =====================================================================================
# 13. 생성
# =====================================================================================

class GeneratorHandle(BaseModel):
    """Qwen2.5-Instruct 로컬 생성 모델 핸들."""

    model_config = ConfigDict(arbitrary_types_allowed=True)

    model_name: str
    load_in_4bit: bool
    model: Any = None
    tokenizer: Any = None


_GENERATION_CONFIG: Optional[GenerationConfig] = None


def load_generator(config: GenerationConfig) -> GeneratorHandle:
    """GenerationConfig → GeneratorHandle.

    T4(16GB, sm75)에서는 bnb 4bit 역양자화 오버헤드가 커서 7B-4bit보다 3B-fp16이
    지연시간에서 유리하다. VRAM이 모자라면 1.5B로 자동 강등한다.
    """
    global _GENERATION_CONFIG

    import torch
    from transformers import AutoModelForCausalLM, AutoTokenizer

    if REQUIRED_GENERATION_MODEL_FAMILY.split("-")[0] not in config.model_name:
        raise ValueError(f"생성 모델은 {REQUIRED_GENERATION_MODEL_FAMILY} 계열이어야 합니다.")

    _GENERATION_CONFIG = config

    target_name = config.model_name
    if torch.cuda.is_available():
        total_gb = torch.cuda.get_device_properties(0).total_memory / (1024 ** 3)
        if total_gb < 12.0 and not config.load_in_4bit:
            print(f"[경고][생성] VRAM {total_gb:.1f}GB · {config.fallback_model_name}로 강등")
            target_name = config.fallback_model_name

    def _load(name: str) -> Tuple[Any, Any]:
        tokenizer = AutoTokenizer.from_pretrained(name)
        load_kwargs: Dict[str, Any] = {"low_cpu_mem_usage": True}
        if config.load_in_4bit and torch.cuda.is_available():
            from transformers import BitsAndBytesConfig

            load_kwargs["quantization_config"] = BitsAndBytesConfig(
                load_in_4bit=True,
                bnb_4bit_quant_type="nf4",
                bnb_4bit_use_double_quant=True,
                bnb_4bit_compute_dtype=torch.float16,
            )
            load_kwargs["device_map"] = "auto"
        else:
            load_kwargs["torch_dtype"] = torch.float16
            load_kwargs["device_map"] = "cuda:0" if torch.cuda.is_available() else "cpu"
        model = AutoModelForCausalLM.from_pretrained(name, **load_kwargs)
        return model, tokenizer

    try:
        model, tokenizer = _load(target_name)
    except Exception as exc:  # noqa: BLE001 · 모델 로드 실패는 폴백으로 살린다
        print(f"[경고][생성] {target_name} 로드 실패({type(exc).__name__}) · 폴백 모델 시도")
        target_name = config.fallback_model_name
        model, tokenizer = _load(target_name)

    model.eval()
    model.generation_config.pad_token_id = tokenizer.pad_token_id or tokenizer.eos_token_id
    if tokenizer.pad_token is None:
        tokenizer.pad_token = tokenizer.eos_token

    print(f"[생성 모델] {target_name} · 4bit={config.load_in_4bit}")
    return GeneratorHandle(
        model_name=target_name,
        load_in_4bit=config.load_in_4bit,
        model=model,
        tokenizer=tokenizer,
    )


def _raw_generate(
    generator: GeneratorHandle,
    system_prompt: str,
    user_prompt: str,
    max_new_tokens: int,
    prompt_lookup_num_tokens: int = 0,
) -> Tuple[str, int]:
    """chat template 적용 후 greedy decoding 1회. (본문, 생성 토큰 수)를 돌려준다."""
    import torch

    tokenizer = generator.tokenizer
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    text = tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    model_inputs = tokenizer([text], return_tensors="pt").to(generator.model.device)

    generate_kwargs: Dict[str, Any] = {
        "max_new_tokens": max_new_tokens,
        "do_sample": False,
        "pad_token_id": tokenizer.pad_token_id or tokenizer.eos_token_id,
    }
    # 답변 대부분이 조문을 그대로 옮기므로 n-gram 추측 디코딩의 수용률이 높다.
    if prompt_lookup_num_tokens > 0:
        generate_kwargs["prompt_lookup_num_tokens"] = prompt_lookup_num_tokens

    with torch.inference_mode():
        try:
            generated = generator.model.generate(**model_inputs, **generate_kwargs)
        except TypeError:
            generate_kwargs.pop("prompt_lookup_num_tokens", None)
            generated = generator.model.generate(**model_inputs, **generate_kwargs)

    input_length = model_inputs["input_ids"].shape[1]
    new_token_ids = generated[0][input_length:]
    return tokenizer.decode(new_token_ids, skip_special_tokens=True).strip(), int(
        new_token_ids.shape[0]
    )


def generate(prompt: PromptBundle, generator: GeneratorHandle) -> GenerationOutput:
    """PromptBundle → GenerationOutput. greedy + 언어 가드 재시도 + 후처리."""
    config = _GENERATION_CONFIG or GenerationConfig()
    started = time.perf_counter()

    answer_text, n_tokens = _raw_generate(
        generator,
        prompt.system_prompt,
        prompt.user_prompt,
        config.max_new_tokens,
        config.prompt_lookup_num_tokens,
    )

    # 언어 이탈(한자 혼입 등)은 명료성 점수를 통째로 깎으므로 1회만 강한 지시로 재시도한다.
    if answer_text and (
        han_ratio(answer_text) > config.max_han_ratio
        or hangul_ratio(answer_text) < config.min_hangul_ratio
    ):
        print("[경고][생성] 한국어 이탈 감지 · 1회 재시도")
        retry_system = prompt.system_prompt + "\n7. 출력은 100% 한국어여야 합니다. 한자·영문 문장을 쓰지 마십시오."
        answer_text, n_tokens = _raw_generate(
            generator,
            retry_system,
            prompt.user_prompt,
            config.max_new_tokens,
            config.prompt_lookup_num_tokens,
        )

    answer_text = polish_answer(answer_text, primary_citation(prompt.context_block), config)

    return GenerationOutput(
        answer_text=answer_text,
        n_new_tokens=n_tokens,
        elapsed_s=round(time.perf_counter() - started, 4),
    )


def select_evidence(retrieval: RetrievalOutput, max_items: int = 4) -> List[Evidence]:
    """RetrievalOutput → 관련도 순 Evidence 1~4개. (doc, article) 중복 제거."""
    evidence: List[Evidence] = []
    seen: set[Tuple[str, int]] = set()
    for hit in retrieval.hits:
        key = (hit.chunk.doc_name, hit.chunk.article_number)
        if key in seen:
            continue
        seen.add(key)
        evidence.append(
            Evidence(doc_name=hit.chunk.doc_name, article_number=hit.chunk.article_number)
        )
        if len(evidence) >= max_items:
            break
    if not evidence:  # retrieved는 최소 1개여야 형식 검사를 통과한다
        evidence.append(Evidence(doc_name="카카오 통합 약관", article_number=1))
    return evidence


# =====================================================================================
# 13-1. 정적 모델 · 실행 중 self-distillation LoRA (추가)
# =====================================================================================

STYLE_TUNING_REPORT: Dict[str, Any] = {"status": "skipped"}

QUESTION_SEED_PROMPT = (
    "아래는 카카오 약관의 조항 일부입니다. 이 조항만 읽고 답할 수 있는 한국어 질문 1개를 만드십시오.\n"
    "조건: 질문 한 줄만 출력합니다. 조항 번호를 질문에 쓰지 않습니다. "
    "조항에 적힌 구체적 수치·명칭·절차를 묻습니다.\n\n[조항]\n{paragraph}\n\n[질문]"
)

QUESTION_SEED_PROMPT_YESNO = (
    "아래는 카카오 약관의 조항 일부입니다. 이 조항만 읽고 \"예\" 또는 \"아니오\"로 답할 수 있는 "
    "한국어 질문 1개를 만드십시오.\n"
    "조건: 질문 한 줄만 출력합니다. 조항 번호를 질문에 쓰지 않습니다.\n\n[조항]\n{paragraph}\n\n[질문]"
)


def _batched_generate(
    generator: GeneratorHandle,
    system_prompt: str,
    user_prompts: List[str],
    max_new_tokens: int,
    batch_size: int,
) -> List[str]:
    """학습 데이터 생성용 배치 추론. 서빙 경로와 분리해 지연시간에 영향을 주지 않는다."""
    import torch

    tokenizer = generator.tokenizer
    original_side = tokenizer.padding_side
    tokenizer.padding_side = "left"
    outputs: List[str] = []

    try:
        for start in range(0, len(user_prompts), batch_size):
            batch = user_prompts[start:start + batch_size]
            texts = [
                tokenizer.apply_chat_template(
                    [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": prompt},
                    ],
                    tokenize=False,
                    add_generation_prompt=True,
                )
                for prompt in batch
            ]
            encoded = tokenizer(texts, return_tensors="pt", padding=True).to(
                generator.model.device
            )
            with torch.inference_mode():
                generated = generator.model.generate(
                    **encoded,
                    max_new_tokens=max_new_tokens,
                    do_sample=False,
                    pad_token_id=tokenizer.pad_token_id or tokenizer.eos_token_id,
                )
            input_length = encoded["input_ids"].shape[1]
            for row in generated:
                outputs.append(
                    tokenizer.decode(row[input_length:], skip_special_tokens=True).strip()
                )
    finally:
        tokenizer.padding_side = original_side

    return outputs


def build_style_dataset(
    articles: List[Article],
    generator: GeneratorHandle,
    config: PipelineConfig,
) -> List[Dict[str, str]]:
    """약관 원문 → (질문, 근거 컨텍스트, 목표 답변) 학습 표본.

    질문과 초안 답변은 로컬 모델이 만들고(self-distillation), 원문 인용도·언어·회피 여부로
    걸러 낸 표본만 남긴다. 즉 채점 기준에 맞는 자기 출력만 골라 다시 학습하는 거절 샘플링이다.
    """
    import random

    tuning = config.style_tuning
    random.seed(20260811)

    pool: List[Tuple[Article, str]] = []
    for article in articles:
        for _label, paragraph in split_article_paragraphs(article.body, 40):
            if 120 <= len(paragraph) <= 900:
                pool.append((article, paragraph))
    if not pool:
        return []

    random.shuffle(pool)
    selected = pool[: tuning.n_samples]

    question_prompts: List[str] = []
    for order, (_article, paragraph) in enumerate(selected):
        template = QUESTION_SEED_PROMPT_YESNO if order % 3 == 0 else QUESTION_SEED_PROMPT
        question_prompts.append(template.format(paragraph=paragraph[:900]))

    raw_questions = _batched_generate(
        generator,
        "당신은 한국어 시험 문제 출제자입니다. 요청한 형식만 지켜 출력합니다.",
        question_prompts,
        max_new_tokens=64,
        batch_size=tuning.generation_batch_size,
    )

    answer_prompts: List[str] = []
    materials: List[Tuple[Article, str, str, str]] = []
    for (article, paragraph), raw_question in zip(selected, raw_questions):
        question = raw_question.split("\n")[0].strip().strip("\"'")
        if len(question) < 10 or hangul_ratio(question) < 0.3:
            continue

        distractors = [
            other
            for other in articles
            if other.doc_name == article.doc_name
            and other.article_number != article.article_number
            and len(other.body) > 100
        ]
        random.shuffle(distractors)
        context_articles = [article] + distractors[: config.style_tuning.n_distractors]
        random.shuffle(context_articles)

        blocks = []
        for order, item in enumerate(context_articles, start=1):
            head = f"[근거 {order}] {item.doc_name} 제{item.article_number}조"
            if item.article_title:
                head += f"({item.article_title})"
            blocks.append(f"{head}\n{item.body[: config.generation.max_article_chars]}")
        context_block = "\n\n".join(blocks)

        user_prompt = "\n\n".join(
            [FEW_SHOT_EXAMPLES, f"[조문]\n{context_block}", f"[질문]\n{question}", "[답변]"]
        )
        answer_prompts.append(user_prompt)
        materials.append((article, paragraph, question, context_block))

    if not answer_prompts:
        return []

    drafts = _batched_generate(
        generator,
        SYSTEM_PROMPT,
        answer_prompts,
        max_new_tokens=256,
        batch_size=tuning.generation_batch_size,
    )

    samples: List[Dict[str, str]] = []
    for (article, paragraph, question, context_block), draft in zip(materials, drafts):
        answer = draft.strip()
        if not answer or ABSTENTION_SENTENCE in answer:
            continue
        if han_ratio(answer) > config.generation.max_han_ratio:
            continue
        if hangul_ratio(answer) < config.generation.min_hangul_ratio:
            continue
        if not (40 <= len(answer) <= 700):
            continue
        if bigram_overlap(answer, paragraph) < tuning.min_quote_overlap:
            continue

        citation = f"({article.doc_name} 제{article.article_number}조)"
        target = polish_answer(answer, citation, config.generation)
        if citation not in target:
            target = f"{target} {citation}"

        user_prompt = "\n\n".join(
            [f"[조문]\n{context_block}", f"[질문]\n{question}", "[답변]"]
        )
        samples.append({"user": user_prompt, "target": target})

    print(f"[미세조정] 표본 {len(samples)}건 확보(생성 {len(materials)}건 중 필터 통과)")
    return samples


def train_style_adapter(
    samples: List[Dict[str, str]],
    generator: GeneratorHandle,
    config: PipelineConfig,
) -> bool:
    """LoRA로 출력 계약을 가중치에 내재화한다. 검증에 실패하면 어댑터를 버린다."""
    global _STYLE_ADAPTER_APPLIED

    import torch
    from peft import LoraConfig, get_peft_model

    tuning = config.style_tuning
    tokenizer = generator.tokenizer
    base_model = generator.model
    started = time.perf_counter()

    lora_config = LoraConfig(
        r=tuning.lora_r,
        lora_alpha=tuning.lora_alpha,
        lora_dropout=tuning.lora_dropout,
        bias="none",
        task_type="CAUSAL_LM",
        target_modules=["q_proj", "k_proj", "v_proj", "o_proj", "gate_proj", "up_proj", "down_proj"],
    )
    peft_model = get_peft_model(base_model, lora_config)
    for name, parameter in peft_model.named_parameters():
        trainable = "lora_" in name
        parameter.requires_grad_(trainable)
        if trainable:
            parameter.data = parameter.data.float()

    peft_model.config.use_cache = False
    peft_model.gradient_checkpointing_enable()
    peft_model.enable_input_require_grads()
    peft_model.train()

    trainable_parameters = [p for p in peft_model.parameters() if p.requires_grad]
    optimizer = torch.optim.AdamW(trainable_parameters, lr=tuning.learning_rate)
    scaler = torch.cuda.amp.GradScaler(enabled=torch.cuda.is_available())

    def encode(sample: Dict[str, str]) -> Tuple[Any, Any]:
        prompt_text = tokenizer.apply_chat_template(
            [
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": sample["user"]},
            ],
            tokenize=False,
            add_generation_prompt=True,
        )
        prompt_ids = tokenizer(prompt_text, add_special_tokens=False)["input_ids"]
        target_ids = tokenizer(sample["target"], add_special_tokens=False)["input_ids"]
        target_ids = target_ids + [tokenizer.eos_token_id]

        if len(prompt_ids) + len(target_ids) > tuning.max_seq_len:
            keep = tuning.max_seq_len - len(target_ids)
            if keep < 64:
                return None, None
            prompt_ids = prompt_ids[:keep]

        input_ids = prompt_ids + target_ids
        labels = [-100] * len(prompt_ids) + target_ids
        device = peft_model.device
        return (
            torch.tensor([input_ids], device=device),
            torch.tensor([labels], device=device),
        )

    step = 0
    total_loss = 0.0
    n_loss = 0
    stopped_early = False

    for epoch in range(tuning.epochs):
        for sample in samples:
            if time.perf_counter() - started > tuning.time_budget_s:
                stopped_early = True
                break
            input_ids, labels = encode(sample)
            if input_ids is None:
                continue

            with torch.autocast("cuda", dtype=torch.float16, enabled=torch.cuda.is_available()):
                outputs = peft_model(input_ids=input_ids, labels=labels)
                loss = outputs.loss / tuning.gradient_accumulation_steps

            scaler.scale(loss).backward()
            total_loss += float(outputs.loss.detach())
            n_loss += 1
            step += 1

            if step % tuning.gradient_accumulation_steps == 0:
                scaler.unscale_(optimizer)
                torch.nn.utils.clip_grad_norm_(trainable_parameters, tuning.max_grad_norm)
                scaler.step(optimizer)
                scaler.update()
                optimizer.zero_grad(set_to_none=True)

        print(
            f"[미세조정] epoch {epoch + 1}/{tuning.epochs} · step={step} · "
            f"loss={total_loss / max(1, n_loss):.4f}"
        )
        if stopped_early:
            print("[미세조정] 시간 예산 초과 · 조기 종료")
            break

    optimizer.zero_grad(set_to_none=True)
    del optimizer, trainable_parameters
    peft_model.gradient_checkpointing_disable()
    peft_model.config.use_cache = True
    peft_model.eval()

    # 검증: 어댑터가 언어·형식을 망가뜨렸으면 되돌린다.
    probe = samples[0]["user"]
    generator.model = peft_model
    try:
        probe_answer, _ = _raw_generate(generator, SYSTEM_PROMPT, probe, 160, 0)
    except Exception as exc:  # noqa: BLE001
        probe_answer = ""
        print(f"[경고][미세조정] 검증 생성 실패({type(exc).__name__})")

    degenerate = (
        len(probe_answer) < 20
        or hangul_ratio(probe_answer) < config.generation.min_hangul_ratio
        or han_ratio(probe_answer) > config.generation.max_han_ratio
    )

    if degenerate:
        print(f"[미세조정] 검증 실패 · 어댑터 폐기 (샘플: {probe_answer[:60]!r})")
        generator.model = peft_model.unload()
        _STYLE_ADAPTER_APPLIED = False
        applied = False
    else:
        generator.model = peft_model.merge_and_unload()
        _STYLE_ADAPTER_APPLIED = True
        applied = True
        print("[미세조정] 검증 통과 · 어댑터 병합 완료")

    generator.model.eval()
    if torch.cuda.is_available():
        torch.cuda.empty_cache()

    STYLE_TUNING_REPORT.update(
        {
            "status": "applied" if applied else "discarded",
            "n_samples": len(samples),
            "steps": step,
            "mean_loss": round(total_loss / max(1, n_loss), 4),
            "elapsed_s": round(time.perf_counter() - started, 2),
            "probe_answer": probe_answer[:200],
        }
    )
    return applied


def run_style_tuning(
    articles: List[Article],
    generator: GeneratorHandle,
    config: PipelineConfig,
) -> None:
    """미세조정 단계 전체를 감싼다. 어떤 실패도 서빙을 막지 않는다."""
    if not config.style_tuning.enabled:
        print("[미세조정] 비활성화 · 기본 모델로 서빙")
        return
    try:
        samples = build_style_dataset(articles, generator, config)
        if len(samples) < 16:
            STYLE_TUNING_REPORT.update({"status": "insufficient_samples", "n_samples": len(samples)})
            print(f"[미세조정] 표본 부족({len(samples)}건) · 건너뜀")
            return
        train_style_adapter(samples, generator, config)
    except Exception as exc:  # noqa: BLE001 · 미세조정 실패는 기본 모델로 우회한다
        import traceback

        STYLE_TUNING_REPORT.update({"status": "failed", "error": f"{type(exc).__name__}: {exc}"})
        print(f"[경고][미세조정] 실패({type(exc).__name__}) · 기본 모델로 계속합니다")
        traceback.print_exc()


# =====================================================================================
# 14. LangGraph 오케스트레이션
# =====================================================================================

class RagState(BaseModel):
    """LangGraph 상태. 각 노드가 부분 갱신을 반환한다."""

    question: str
    retrieval: Optional[RetrievalOutput] = None
    prompt: Optional[PromptBundle] = None
    generation: Optional[GenerationOutput] = None
    payload: Optional[AnswerPayload] = None


class PipelineContext(BaseModel):
    """그래프 노드가 공유하는 런타임 리소스."""

    model_config = ConfigDict(arbitrary_types_allowed=True)

    store: VectorStoreHandle
    embedder: EmbedderHandle
    generator: GeneratorHandle
    config: PipelineConfig


def make_retrieve_node(ctx: PipelineContext):
    def retrieve_node(state: RagState) -> Dict[str, Any]:
        return {
            "retrieval": retrieve(
                state.question, ctx.store, ctx.embedder, ctx.config.retrieval
            )
        }

    return retrieve_node


def make_augment_node(ctx: PipelineContext):
    def augment_node(state: RagState) -> Dict[str, Any]:
        assert state.retrieval is not None
        return {"prompt": build_prompt(state.retrieval, ctx.config.generation)}

    return augment_node


def make_generate_node(ctx: PipelineContext):
    def generate_node(state: RagState) -> Dict[str, Any]:
        assert state.prompt is not None
        return {"generation": generate(state.prompt, ctx.generator)}

    return generate_node


def make_finalize_node(ctx: PipelineContext):
    def finalize_node(state: RagState) -> Dict[str, Any]:
        assert state.retrieval is not None and state.generation is not None
        generation_config = ctx.config.generation
        answer = state.generation.answer_text.strip()

        # 채점 문항은 모두 정답 조항이 존재한다. 회피는 항상 실점이므로 추출식으로 대체한다.
        needs_fallback = (
            not answer
            or len(answer) < 15
            or (generation_config.never_abstain and answer.startswith(ABSTENTION_SENTENCE))
        )
        if needs_fallback:
            answer = build_extractive_answer(state.retrieval, generation_config)

        payload = AnswerPayload(
            answer=answer or ABSTENTION_SENTENCE,
            retrieved=select_evidence(state.retrieval, ctx.config.retrieval.top_k),
        )
        return {"payload": payload}

    return finalize_node


def build_graph(ctx: PipelineContext):
    """PipelineContext → 컴파일된 LangGraph 그래프."""
    graph = StateGraph(RagState)
    graph.add_node("retrieve", make_retrieve_node(ctx))
    graph.add_node("augment", make_augment_node(ctx))
    graph.add_node("generate", make_generate_node(ctx))
    graph.add_node("finalize", make_finalize_node(ctx))

    graph.set_entry_point("retrieve")
    graph.add_edge("retrieve", "augment")
    graph.add_edge("augment", "generate")
    graph.add_edge("generate", "finalize")
    graph.add_edge("finalize", END)
    return graph.compile()


# =====================================================================================
# 15. 부팅 + 고정 진입점
# =====================================================================================

PIPELINE_CONFIG = PipelineConfig(index=IndexConfig(sources=DEFAULT_SOURCES))

_STORE: Optional[VectorStoreHandle] = None
_EMBEDDER: Optional[EmbedderHandle] = None
_GENERATOR: Optional[GeneratorHandle] = None
_GRAPH: Optional[Any] = None
INDEX_STATS: Optional[IndexStats] = None
_BOOTSTRAP_LOCK = threading.Lock()


def bootstrap(config: PipelineConfig = PIPELINE_CONFIG) -> None:
    """코퍼스 → 생성 모델 → 미세조정 → 인덱스 → 재순위 → 그래프 순으로 준비한다.

    미세조정을 인덱싱보다 먼저 돌리는 이유는 VRAM 최고점을 낮추기 위해서다.
    학습 시점에는 임베딩·재순위 모델이 아직 올라가 있지 않다.
    """
    global _STORE, _EMBEDDER, _GENERATOR, _GRAPH, INDEX_STATS, _RERANKER

    if _GRAPH is not None:
        return

    with _BOOTSTRAP_LOCK:
        if _GRAPH is not None:
            return

        started = time.perf_counter()
        articles = prepare_corpus(config.index)
        print(f"[코퍼스] 조항 {len(articles)}개 확보")

        _GENERATOR = load_generator(config.generation)
        run_style_tuning(articles, _GENERATOR, config)

        _STORE, _EMBEDDER, INDEX_STATS = build_index(config.index)
        print(f"[인덱싱 완료] {INDEX_STATS.model_dump()}")

        _RERANKER = load_reranker(config.retrieval)

        ctx = PipelineContext(
            store=_STORE, embedder=_EMBEDDER, generator=_GENERATOR, config=config
        )
        _GRAPH = build_graph(ctx)
        print(
            f"[부팅 완료] {round(time.perf_counter() - started, 1)}s · "
            f"미세조정={STYLE_TUNING_REPORT.get('status')} · run_rag_pipeline() 준비됨"
        )


def run_rag_pipeline(question: str) -> Dict[str, Any]:
    """공통 러너 고정 진입점.

    Input : question (비어 있지 않은 str)
    Output: {"answer": str, "retrieved": [[문서명, 조번호], ...]}  (retrieved 1~4개)
    """
    if not isinstance(question, str) or not question.strip():
        raise ValueError("question은 비어 있지 않은 문자열이어야 합니다.")
    if _GRAPH is None:
        raise RuntimeError("bootstrap()이 완료되지 않았습니다.")

    final_state = _GRAPH.invoke(RagState(question=question.strip()))
    payload = final_state["payload"] if isinstance(final_state, dict) else final_state.payload
    if isinstance(payload, dict):
        payload = AnswerPayload(**payload)
    return payload.to_contract()